from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "database" / "isms.sqlite"
OUT_PATH = ROOT / "reports" / "ISMS_MediNova_UTM_2026.docx"


def query(sql, params=()):
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        return [dict(row) for row in conn.execute(sql, params).fetchall()]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True, size=font_size)
        set_cell_shading(header_cells[index], "EDEFF2")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
      run.font.name = "Arial"
      run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(30, 30, 30)


def set_margins(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "TSI · ISO 27001 / ISO 27005 · MediNova Clinic"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def build():
    org = query("SELECT * FROM organizations LIMIT 1")[0]
    departments = query("SELECT name, description FROM departments ORDER BY id")
    assets = query("SELECT * FROM assets ORDER BY id")
    threats = query("SELECT threats.*, assets.name AS asset_name, assets.category AS asset_category FROM threats JOIN assets ON assets.id = threats.asset_id ORDER BY threats.id")
    risks = query("SELECT risks.*, assets.name AS asset_name FROM risks JOIN assets ON assets.id = risks.asset_id ORDER BY risk_score DESC")
    quantitative = query("SELECT * FROM quantitative_risks ORDER BY ale DESC")
    controls = query("SELECT * FROM controls ORDER BY control_type, control_function, id")
    policies = query("SELECT * FROM policies ORDER BY id")

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    set_margins(section)

    # Title page without running header.
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Технический университет Молдовы (UTM)")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Факультет информационных технологий")
    r.font.name = "Arial"
    r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Лабораторная работа")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Система менеджмента информационной безопасности на основе ISO 27001 и ISO 27005")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    meta = [
        ("Предмет", "TSI"),
        ("Группа", "TI-242FR"),
        ("Выполнили", "Alina Bocearova, Elepin Vladislav, Igor Tiora"),
        ("Организация проекта", "ООО «MediNova Clinic»"),
        ("Год", "2026"),
    ]
    doc.add_paragraph()
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        set_cell_text(table.rows[i].cells[0], label, bold=True, size=11)
        set_cell_text(table.rows[i].cells[1], value, size=11)
        set_cell_shading(table.rows[i].cells[0], "F2F2F2")

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Кишинёв, 2026")
    r.font.name = "Arial"
    r.font.size = Pt(12)

    doc.add_section(WD_SECTION.NEW_PAGE)
    content_section = doc.sections[-1]
    set_margins(content_section)
    add_header_footer(content_section)

    add_heading(doc, "Lab Report - Document Enterprise Cybersecurity Issues", 1)
    add_body(doc, "Отчёт оформлен по структуре лабораторной работы: Objectives, Scenario, Instructions, Part 1-4, таблицы активов, угроз и мер снижения риска, а также элементы ISO 27001 и ISO 27005.")

    add_heading(doc, "Objectives", 1)
    for item in [
        "Part 1: Зафиксировать оценку проблем кибербезопасности выбранной организации.",
        "Part 2: Определить основные типы активов, которыми владеет организация.",
        "Part 3: Перечислить угрозы для каждого типа активов.",
        "Part 4: Предложить методы снижения риска для каждой угрозы.",
        "Дополнительно: показать область применения ISMS, реестр рисков, контроли, политики и итоговый план внедрения.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Scenario", 1)
    add_body(doc, f"ООО «MediNova Clinic» — частная медицинская клиника и диагностический центр. В клинике работает около {org['employees']} сотрудников. Организация обрабатывает персональные данные пациентов, электронные медицинские карты, лабораторные результаты, финансовую и страховую документацию.")
    add_body(doc, "Клиника должна обеспечивать конфиденциальность медицинских и персональных данных, целостность медицинских записей и доступность критических сервисов.")

    add_heading(doc, "Instructions", 1)
    for item in [
        "Information/Data Assets — данные, которые используются, хранятся или передаются организацией.",
        "Software Assets — операционные системы, серверное ПО, базы данных и корпоративные приложения.",
        "Physical Assets — серверы, рабочие станции, камеры, считыватели и другое оборудование.",
        "Network Assets — сети, сетевые сервисы, VPN, Wi-Fi, firewall и интернет-канал.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Part 1: Record the assessment of cybersecurity issues", 1)
    add_heading(doc, "Описание организации", 2)
    add_body(doc, org["description"])
    add_body(doc, f"Область применения ISMS: {org['scope']}")
    add_body(doc, f"Цели безопасности: {org['security_objectives']}")
    add_heading(doc, "Подразделения", 2)
    add_table(doc, ["Подразделение", "Описание"], [[d["name"], d["description"]] for d in departments], font_size=9)

    add_heading(doc, "Основные проблемы кибербезопасности", 2)
    add_table(doc, ["Проблема", "Описание", "Возможное влияние"], [
        ["Защита медицинских данных", "Клиника хранит персональные данные пациентов, электронные медицинские карты и лабораторные результаты.", "Утечка данных, нарушение конфиденциальности, репутационный ущерб."],
        ["Доступность медицинских сервисов", "Регистрация пациентов, запись на прием и доступ врачей к медицинским данным зависят от IT-инфраструктуры.", "Простой клиники, задержка медицинских услуг, финансовые потери."],
        ["Удаленный доступ", "Врачи и администраторы используют VPN и удаленный доступ к внутренним системам.", "Компрометация учетных записей, несанкционированный доступ."],
        ["Резервное копирование", "Медицинские данные должны быть восстановимы после ransomware, сбоя или ошибки.", "Потеря данных и невозможность продолжать работу."],
        ["Физическая безопасность", "Серверы, CCTV и система контроля доступа должны быть защищены от физического воздействия.", "Кража оборудования, повреждение инфраструктуры, нарушение доступности."],
    ], font_size=8)

    add_heading(doc, "Part 2: Record the different types of assets", 1)
    categories = ["Информация / данные", "Программное обеспечение", "Аппаратное обеспечение", "Сеть"]
    for category in categories:
        add_heading(doc, category, 2)
        rows = [
            [a["name"], a["owner_department"], a["confidentiality_level"], a["integrity_level"], a["availability_level"], a["business_value"]]
            for a in assets if a["category"] == category
        ]
        add_table(doc, ["Актив", "Владелец", "К", "Ц", "Д", "Ценность"], rows, font_size=8)

    add_heading(doc, "Part 3: List the threats for each asset type", 1)
    add_heading(doc, "Вопрос: в чем разница между угрозой и уязвимостью?", 2)
    add_body(doc, "Уязвимость — это слабое место актива, процесса или контроля, которое может быть использовано. Угроза — это возможное событие или действие, которое использует уязвимость и приводит к ущербу.")

    for category in categories:
        add_heading(doc, category, 2)
        category_assets = [a for a in assets if a["category"] == category]
        rows = []
        for asset in category_assets:
            related = [t for t in threats if t["asset_id"] == asset["id"]]
            if related:
                threat_text = "\n".join([f"{t['vulnerability']} -> {t['threat']} -> {t['consequence']}" for t in related])
                mitigation = "\n".join([t["mitigation"] for t in related])
            else:
                threat_text = "Риск не выявлен в учебном наборе данных"
                mitigation = "Периодический пересмотр и базовые меры защиты"
            rows.append([asset["name"], threat_text, mitigation])
        add_table(doc, ["Активы", "Угрозы", "Меры снижения риска"], rows, font_size=8)

    add_heading(doc, "Part 4: Recommend mitigation techniques", 1)
    add_body(doc, "Для снижения рисков предложены физические, технические и административные контроли. Они разделяются по функции: предупреждающие, обнаруживающие и корректирующие.")
    add_table(doc, ["Тип", "Функция", "Контроль", "Статус", "Ответственный отдел"], [
        [c["control_type"], c["control_function"], c["name"], c["implementation_status"], c["responsible_department"]]
        for c in controls
    ], font_size=7)

    add_heading(doc, "ISO 27005 Risk Assessment", 1)
    add_body(doc, "Для оценки риска используется учебная модель TP/VL/SEV/DET. TP — вероятность угрозы, VL — уровень уязвимости, SEV — влияние на бизнес, DET — возможность обнаружения.")
    add_body(doc, "Формула: risk = round((((TP - 1) + (VL - 1) + (SEV - 1) + (5 - DET)) / 16) × 100).")
    add_table(doc, ["ID", "Актив", "Уязвимость", "Угроза", "TP", "VL", "SEV", "DET", "Риск %", "Уровень"], [
        [r["id"], r["asset_name"], r["vulnerability"], r["threat"], r["tp"], r["vl"], r["sev"], r["det"], r["risk_score"], r["risk_level"]]
        for r in risks
    ], font_size=7)

    add_heading(doc, "Quantitative Risk Analysis", 1)
    total_ale = sum(float(q["ale"] or 0) for q in quantitative)
    add_body(doc, f"SLE = Asset Value × Exposure Factor. ALE = SLE × ARO. Общий ожидаемый годовой ущерб по учебным сценариям: {total_ale:.2f} EUR.")
    add_table(doc, ["Актив", "Угроза", "Стоимость", "EF", "SLE", "ARO", "ALE"], [
        [q["asset_name"], q["threat"], q["asset_value"], q["exposure_factor"], q["sle"], q["aro"], q["ale"]]
        for q in quantitative
    ], font_size=8)

    add_heading(doc, "Security Policies", 1)
    add_table(doc, ["Политика", "Цель", "Периодичность"], [
        [p["title"], p["purpose"], p["review_frequency"]]
        for p in policies
    ], font_size=8)

    add_heading(doc, "Reflection", 1)
    reflections = [
        ("Почему полезно категоризировать активы при определении угроз и мер защиты?", "Категоризация помогает структурировать анализ. Для данных, программного обеспечения, оборудования и сетей характерны разные угрозы и разные меры защиты."),
        ("Могут ли разные угрозы иметь одинаковые или похожие меры снижения риска?", "Да. Например, MFA снижает риск кражи учетных данных для VPN, почты и внутренних систем. Резервное копирование помогает при ransomware, ошибках персонала и сбоях оборудования."),
        ("Что показывает применение знаний о киберугрозах к смоделированной организации?", "Проект показывает, что защита организации требует комплекса технических, административных и физических мер, регулярной оценки рисков, обучения персонала, мониторинга и реагирования на инциденты."),
    ]
    for question, answer in reflections:
        add_heading(doc, question, 2)
        add_body(doc, answer)

    add_heading(doc, "Conclusion", 1)
    add_body(doc, "В результате была разработана учебная модель ISMS для ООО «MediNova Clinic». Отчет демонстрирует полный цикл анализа: описание организации, определение активов, угроз и уязвимостей, оценку рисков по ISO 27005, выбор контролей, разработку политик и формирование рекомендаций.")

    doc.core_properties.title = "Система менеджмента информационной безопасности для ООО «MediNova Clinic»"
    doc.core_properties.author = "Alina Bocearova, Elepin Vladislav, Igor Tiora"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
